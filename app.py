import streamlit as st
from openai import OpenAI
from io import BytesIO
import docx
from st_social_media_links import SocialMediaIcons
import os
from dotenv import load_dotenv

# Configurar la página
st.set_page_config(page_title="Contenido para Ciencia de Datos con GPT-4o", page_icon="🤖")


# Función genérica para solicitudes a OpenAI
def generate_content(task, prompt, client, max_tokens=1500, temperature=0.7):
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": task},
                {"role": "user", "content": prompt}
            ],
            max_tokens=max_tokens,
            temperature=temperature
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Ocurrió un error: {str(e)}")
        return None


# Función para crear archivos Word
def create_word_doc(content, title="Contenido Generado"):
    doc = docx.Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Configurar el logo y título
st.image("https://cdn-icons-png.flaticon.com/512/4824/4824797.png", width=80)
st.markdown("<h1 style='text-align: center;'>SMART DATA SCIENCE</h1>", unsafe_allow_html=True)

# Configurar la barra lateral para navegación
section = st.sidebar.selectbox("Selecciona una opción", ["Resúmenes Técnicos", "Código Python", "Dataset"])


def prompt_for_api():
    """Pide la clave de API al usuario."""
    st.session_state.api_key = st.text_input("Por favor ingresa tu clave de OpenAI API:", type="password")
    if st.session_state.api_key:
        try:
            client = OpenAI(api_key=st.session_state.api_key)
            st.success("Clave aceptada. Ahora puedes interactuar con la aplicación.")
            return client
        except Exception as e:
            st.error(f"No se pudo conectar con la clave API: {str(e)}")
            return None
    return None


# Lógica principal para cada sección
if section == "Resúmenes Técnicos":
    st.markdown("### Generación de Resúmenes Técnicos")
    st.markdown(
        """
        🔍 **Ejemplos de uso**:
    - (Análisis de clustering, Redes neuronales profundas, Visualización de datos, Optimización de algoritmos)
        """
    )
    if st.button("Usar Resúmenes Técnicos"):
        client = prompt_for_api()
        if client:  # Si la clave es válida, permitir la interacción
            topic = st.text_input("Ingrese un tema de ciencia de datos:")
            if st.button("Generar Resumen"):
                if topic:
                    task = "Eres un experto en ciencia de datos. Proporciona resúmenes técnicos claros con ejemplos prácticos."
                    summary = generate_content(task, f"Explica: {topic}", client, temperature=0.6)
                    if summary:
                        st.markdown(f"### Resumen sobre {topic}")
                        st.markdown(summary)
                        st.download_button(
                            "Descargar Resumen",
                            data=create_word_doc(summary, title=f"Resumen sobre {topic}"),
                            file_name="resumen_tecnico.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

elif section == "Código Python":
    st.markdown("### Generación de Código")
    st.markdown(
        """
        💡 **Ejemplos de uso**:
        - (Creación de gráficos, Conectar a bases de datos SQL, Web scraping de datos de precios, Crear una API REST básica con FastAPI)
        """
    )
    if st.button("Usar Generador de Código"):
        client = prompt_for_api()
        if client:
            description = st.text_input("Describe qué código necesitas:")
            if st.button("Generar Código"):
                if description:
                    task = "Eres un experto en Python. Genera código eficiente sin explicaciones adicionales."
                    code = generate_content(task, f"Escribe código para: {description}", client, temperature=0.2)
                    if code:
                        st.code(code, language="python")
                        st.download_button("Descargar Código", data=code, file_name="codigo.py", mime="text/plain")

elif section == "Dataset":
    st.markdown("### Generación de Dataset")
    st.markdown(
        """
        📊 **Ejemplos de uso**:
- (Datos de ventas para 2024, Clientes de e-commerce, Datos demográficos de una región, Precios de criptomonedas históricas)
        """
    )
    if st.button("Usar Generador de Dataset"):
        client = prompt_for_api()
        if client:
            description = st.text_input("Describe el dataset que necesitas:")
            if st.button("Generar Dataset"):
                if description:
                    task = "Genera datos en formato CSV con encabezados y al menos 100 filas."
                    data = generate_content(task, f"Dataset sobre: {description}", client, temperature=0.5)
                    if data:
                        st.download_button(
                            "Descargar Dataset",
                            data=data.encode('utf-8'),
                            file_name="dataset.csv",
                            mime="text/csv"
                        )

# Pie de página
st.markdown(
    """
    <div style="text-align: center;">
        <strong>Desarrollador:</strong> Edwin Quintero Alzate | <strong>Email:</strong> egqa1975@gmail.com
    </div>
    """,
    unsafe_allow_html=True
)
SocialMediaIcons([
    "https://www.facebook.com/edwin.quinteroalzate",
    "https://www.linkedin.com/in/edwinquintero0329/",
    "https://github.com/Edwin1719",
]).render()
